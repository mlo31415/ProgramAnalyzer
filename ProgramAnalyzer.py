from __future__ import annotations

import re
from collections import defaultdict

import json
import os.path
import difflib
import re as RegEx
from datetime import datetime
import csv

import docx
from docx.shared import Pt, Inches
from docx import text
from docx.text import paragraph
from docx.enum.section import WD_ORIENTATION

import openpyxl

import numpy as np
from googleapiclient.discovery import build
from google.oauth2 import service_account
from googleapiclient.errors import HttpError

from HelpersPackage import PyiResourcePath, ParmDict, ReadListAsParmDict, MessageLog, SquareUpMatrix, RemoveEmptyRowsFromMatrix
from HelpersPackage import GetParmFromParmDict, SearchAndReplace, UnicodeToHtml

from ScheduleElement import ScheduleElement
from Item import Item
from Person import Person
from Log import Log, LogClose, LogError
from NumericTime import NumericTime


def main():
    # *************************************************************************************************
    # *************************************************************************************************
    # MAIN
    # Read and analyze the spreadsheet

    Log("Started")

    # Read the parameters.
    # This includes the names of the specific tabs to be used.
    parms=ReadListAsParmDict('parameters.txt')
    if parms is None or len(parms) == 0:
        MessageLog(f"Can't open/read {os.getcwd()}/parameters.txt\nProgramAnalyzer terminated.")
        exit(999)

    # Create the reports subfolder if none exists
    reportsdir=GetParmFromParmDict(parms, "reportsdir", "Reports")
    if not os.path.exists(reportsdir):
        os.mkdir(reportsdir)
        Log(f"Reports directory {os.getcwd()}/{reportsdir} created ")

    # Are we getting the program from Google docs or from an Excel spreadsheet?
    source=GetParmFromParmDict(parms, "source", "Google")
    if source.lower() == "google":

        Log("Loading program from Google docs")
        with open(GetParmFromParmDict(parms, "credentials")) as source:
            info=json.load(source)
            Log("Json read")

        if info is None:
            MessageLog("Json file is empty")
            exit(999)
        Log("Spreadsheet credentials read")

        credentials=service_account.Credentials.from_service_account_info(info)
        Log("Credentials established", Flush=True)

        service=build('sheets', 'v4', credentials=credentials)
        Log("Service established", Flush=True)

        # Call the Sheets API to load the various tabs of the spreadsheet
        googleSheets=service.spreadsheets()
        SPREADSHEET_ID=GetParmFromParmDict(parms, "SheetID")  # This is the ID of the specific spreadsheet we're reading
        scheduleCells=ReadSheetFromGoogleTab(googleSheets, SPREADSHEET_ID, parms, "ScheduleTab")
        precisCells=ReadSheetFromGoogleTab(googleSheets, SPREADSHEET_ID, parms, "PrecisTab")
        peopleCells=ReadSheetFromGoogleTab(googleSheets, SPREADSHEET_ID, parms, "PeopleTab")
        parameterCells=ReadSheetFromGoogleTab(googleSheets, SPREADSHEET_ID, parms, "ControlTab")

    else:
        Log(f"Loading program from '{source}'")

        workbook=openpyxl.load_workbook(source)
        scheduleCells=ReadSheetFromXLSXTab(workbook, parms, "ScheduleTab")
        precisCells=ReadSheetFromXLSXTab(workbook, parms, "PrecisTab")
        peopleCells=ReadSheetFromXLSXTab(workbook, parms, "PeopleTab")
        parameterCells=ReadSheetFromXLSXTab(workbook, parms, "ControlTab")

    # Read parameters from the Control sheet
    startingDay="Friday"
    for row in parameterCells:
        if row[0] == "Starting day":
            if len(row) > 1:
                startingDay=row[1].strip()
                startingDay=startingDay[0].upper()+startingDay.lower()[1:]  # Force the capitalization to be right

    # Reorganize the dayList so it starts with our starting day. It's extra-long so that clipping days from the front will still leave a full week.
    if not NumericTime.SetStartingDay(startingDay):
        LogError("Can't interpret ControlTab:Starting day='"+startingDay+"'.  Will use 'Friday'")
        NumericTime().SetStartingDay("Friday")

    # We're done with reading the spreadsheet. Now analyze the data.
    # ******
    # Start reading ths spreadsheet and building the participants and items databases (dictionaries)
    # Note that time and room are redundant and could be pulled out of the Items dictionary
    gItems: dict[str, Item]={}  # A dictionary keyed by item name containing an Item (time, room, people-list, moderator), where people-list is the list of people on the item
    gTimes: list[NumericTime]=[]  # A list of times found in the spreadsheet.
    gPersons: defaultdict[str, Person]=defaultdict(Person)   # A dict of Persons keyed by the people key (full name)
    gRoomNames: list[str]=[]    # The list of room names corresponding to the columns in gItems


    #***********************************************************************
    # Analyze the People cells and add the information to gPersons

    # Start by removing empty rows and padding all rows out to make the array rectangular
    peopleCells=SquareUpMatrix(RemoveEmptyRowsFromMatrix(peopleCells))
    columnLabels=peopleCells[0]

    # Check for duplicate column headers -- this is always a fatal error.
    for item in columnLabels:
        if columnLabels.count(item) != 1:
            MessageLog(f"'{item}' appears other than once as a column header.  Terminating.")
            exit(999)

    # Now read the remaining rows one by one, storing the cells in a ParmDict with the column header as key.
    for irow, row in enumerate(peopleCells[1:]):
        pd=ParmDict(CaseInsensitiveCompare=True)
        for i, val in enumerate(row):
            pd[columnLabels[i]]=val

        # Now, we need to form a Fullname for the Person.
        # If there is a Fullname column, use that.
        fullname=""
        if pd.Exists("full name") and pd["full name"] != "":
            fullname=pd["full name"]
        else:
            # Create a fullname out of fname+lname
            # Got to handle the case where one or the other name is missing or empty
            if pd.Exists("fname"):
                fullname=pd["fname"].strip()
            if pd.Exists("lname"):
                fullname=(fullname+" "+pd["lname"].strip()).strip()

        if fullname == "":
            LogError(f"*** Can't find or construct a non-null full name for row {irow+1}")
            LogError("      Col Names: "+str(columnLabels))
            LogError("      Row Data:  "+str(row))
            continue

        pd["Fullname"]=fullname
        gPersons[fullname]=Person(fullname, pd)       # Store the email and response in a Person structure indexed by the full name




    # Now process the main schedule row by row
    # When we find a row with data in column 0, we have found a new time. This is a time row.
    # A time row contains items.
    # A time row will normally be followed by a people row containing the participants for those items.
    # A people row does *not* have content in column 0.

    # The rows for a particular time con be a single row or two rows, in which case the 2nd row contains the people scheduled on that item.
    # Rows that are blank or start with a # as the 1st character of column 0 are ignored
    # Compress out the ignored rows
    cleanedSchedualCells: list[list[str]]=[]
    for row in scheduleCells:
        if len(row) == 0:  # Ignore empty rows
            continue
        # Skip rows where the first character in the row is a "#"
        s="".join([r.strip() for r in row])
        if s[0] == "#":
            continue
        # Cells which start with a "#" are treated as blank
        for cell in row:
            if cell.strip().startswith("#"):
                cell=""
        cleanedSchedualCells.append(row)

    cleanedSchedualCells=SquareUpMatrix(cleanedSchedualCells)

    # Now compress out non-room and non-time columns
    # This will leave one time column on the left followed by all the room columns
    # We will drop columns even if they have something in them if they are not headed by a room name
    # We wprk in the transposed cleanedSchedualCells, since it's much easier to delete rows than columns
    temp=np.array(cleanedSchedualCells).T.tolist()  # Use numpy to transpose the array
    cleanedSchedualCells=[temp[0]]    # Copy over the time row
    for row in temp[1:]:
        if len(row[0].strip()) > 0:     # Copy over any rows with text in the first cell
            cleanedSchedualCells.append(row)
    cleanedSchedualCells=np.array(cleanedSchedualCells).T.tolist()  # And transpose it back

    # Move the room names line out of cleanedSchedualCells and into gRoomNames
    gRoomNames=[r.strip() for r in cleanedSchedualCells[0]] # Get the room names which are in the first row of the scheduleCells tab
    if len(gRoomNames) == 0:
        LogError("Room names line (1st row of the schedule tab) is blank.")
        return
    cleanedSchedualCells=cleanedSchedualCells[1:]


    # Now we have just the schedule rows.  They are of two types:
    #       A time/items row, which contains a time in column 0 and may contain items in some or all of the rest of the columns
    #       A people row which follows a time row and has column 0 empty. This may contain a list of people for each of the items
    # Process them.
    rowIndex=0
    while rowIndex < len(cleanedSchedualCells):
        # The first row must be a time/items row.
        row=cleanedSchedualCells[rowIndex]
        if len(row[0]) == 0:     # Time/items rows have content in the 1st column. Is it a time/items row?
            LogError("Error reading schedule tab: The row below is a people row; we were expecting a time/items row:")
            LogError("       row="+" ".join(row))
            rowIndex+=1
            continue
        rowFirst=row
        rowIndex+=1

        # Possibly followed by a people row
        rowSecond=None
        if rowIndex < len(cleanedSchedualCells):
            row=cleanedSchedualCells[rowIndex]   # Peek ahead to the next row
            if len(row[0]) == 0:
                # We found a people row
                rowSecond=row
                rowIndex+=1

        # Get the time from rowFirst and add it to gTimes
        time=NumericTime(rowFirst[0])
        if time not in gTimes:
            gTimes.append(time)  # We want to allow duplicate time rows, just-in-case

        # Looking at the rest of the row, there may be text in one or more of the room columns which defines an item
        for col, roomName in enumerate(gRoomNames):
            if col == 0:    # Time is in col 0, so we don't want to look at that
                continue

            # This has to be an item name since it's a cell containing text in a row that starts with a time and in a column that starts with a room
            itemName=rowFirst[col].strip()
            if len(itemName) > 0 and not itemName.startswith("#"):  # It is only an item if the cell contains text

                # In some cases, the item may have a generic name, e.g.,  "Reading", "Autographs".  This name will be used in multiple places, but
                # We require a unique name to track the isons of people with items.  If an item name is already in gItems, we uniquify the next use of that item name
                # by appending rom/day/time to it.
                # Note that anything in {curly brackets} is ignored when printing, etc.
                lst, val=SearchAndReplace("(<.*?>)", itemName, "")
                itemNameStripped=val.strip()
                if itemNameStripped in gItems:
                    itemName+=" {"+roomName+" "+str(time)+"}"
                    Log(f"Item Name decorated {itemName}")

                # Was there a people row following this time/items row?
                if rowSecond is not None:
                    # We indicate items which go for an hour, but have some people in one part and some in another using a special notation in the people list.
                    # Robert A. Heinlein, [0.5] John W. Campbell puts RAH on the hour and JWC a half-hour later.
                    # There is much messiness in this.
                    # We look for the [##] in the people list.  If we find it, we divide the people list in half and create two items with separate plists.
                    r=RegEx.match(r"(.*)\[([0-9.]*)](.*)", rowSecond[col])
                    if r is None:
                        AddItemWithPeople(gItems, time, roomName, itemName, rowSecond[col])
                    else:
                        # Sometimes the first person can have a trailing comma, e.g., Socrates, [0.0] Plato.  Drop it.
                        plist1=r.groups()[0].strip().removesuffix(",")
                        deltaT=float(r.groups()[1].strip())
                        plist2=r.groups()[2].strip()
                        AddItemWithPeople(gItems, time, roomName, itemName, plist1, length=deltaT)
                        newTime=time+deltaT
                        if newTime not in gTimes:
                            gTimes.append(newTime)
                        # This second instance will need to have a distinct item name, so add {#2} to the item name
                        AddItemWithPeople(gItems, newTime, roomName, itemName+" {#2}", plist2, length=1.0-deltaT)   #TODO: Do we want to handle divisions other thin into 1/2?
                else:  # We have an item with no people on it.
                    AddItemWithoutPeople(gItems, time, roomName, itemName, 1.0)


    # Extract information from Items, etc., to be used to process schedules
    gSchedules: dict[str, list[ScheduleElement]]=defaultdict(list)  # A dictionary keyed by a person's name containing a ScheduleElement list
    # ScheduleElement is the (time, room, item, moderator) tuples, of an item that that person is on.

    # Used so that the gSchedules XML contains entries for unscheduled people which will be used in ProgramMailAnalyzer to handle things like invitations
    for person in gPersons:
        gSchedules[person]=[ScheduleElement(PersonName=person, IsDummy=True, )]

    for item in gItems.values():
        for personName in item.People:  # For each person listed on this item
            ismod, personName=CheckModFlag(personName)
            gSchedules[personName].append(ScheduleElement(PersonName=personName, Time=item.Time, Length=item.Length, Room=item.Room, ItemName=item.Name, IsMod=ismod))  # And append a tuple with the time, room, item name, and moderator flag

    # Make sure times are sorted into ascending order.
    # The simple sort works because the times are stored as numeric hours since start of first day.
    gTimes.sort()

    # Create a timestemp
    timestamp=f"Generated: {datetime.now():%A %B %d, %Y at %H:%M:%S}\n\n"

    #***********************************************************************
    # Analyze the Precis cells and add the information to gItems
    # The first row is column labels. So ignore it.
    precisCells=precisCells[1:]

    # The rest of the rows of the tab contains the title in the first column and the precis in the second
    count: int=0
    fname=os.path.join(reportsdir, "Diag - precis without items.txt")
    with open(fname, "w") as f:
        print("Precis without corresponding items:", file=f)
        print(timestamp,  file=f)
        for row in precisCells:
            row=[r.strip() for r in row]    # Get rid of leading and trailing blanks
            if len(row) > 1 and len(row[0]) > 0 and len(row[1]) > 0: # If both the item name and the precis exist, store them in the precis table.
                itemname=row[0]
                if itemname in gItems:
                    gItems[itemname].Precis=row[1]
                else:
                    count+=1
                    print("   "+itemname, file=f)
        if count == 0:
            print("    None found", file=f)


    #*************************************************************************************************
    #*************************************************************************************************
    # Generate reports
    # The first reports are all error reports or checking reports


    # We have precois which include material in ((double parens).  This material goes into some reports, but not all.
    # Strip the non-public stuff -- ((in double parens)) from one precis
    def ScrubPrecis(pre: str) -> str:
        return re.sub(r"\(\(.*\)\)", "", pre, flags=re.DOTALL)


    #******
    # Check for people in the schedule who are not in the people tab
    fname=os.path.join(reportsdir, "Diag - People in schedule but not in People.txt")
    with open(fname, "w") as f:
        print("People who are scheduled but not in People:", file=f)
        print("(Note that these may be due to spelling differences, use of initials, etc.)", file=f)
        print(timestamp,  file=f)
        count=0
        for personname in gSchedules.keys():
            if personname not in gPersons.keys():
                count+=1
                print("   "+personname, file=f)
        if count == 0:
            print("    None found", file=f)


    #******
    # Check for people in the schedule whose response is not 'y'
    fname=os.path.join(reportsdir, "Diag - People in schedule and in People but whose response is not 'y'.txt")
    with open(fname, "w") as f:
        print("People who are scheduled and in People but whose response is not 'y':", file=f)
        print(timestamp,  file=f)
        count=0
        for personname in gSchedules.keys():
            if any([not x.IsDummy for x in gSchedules[personname]]):
                if personname in gPersons.keys():
                    if not gPersons[personname].RespondedYes:
                        count+=1
                        print(f"   {personname} has a response of '{gPersons[personname].Response}'", file=f)
        if count == 0:
            print("    None found", file=f)


    #******
    # Check for people with bogus email addresses
    fname=os.path.join(reportsdir, "Diag - People with suspect email addresses.txt")
    with open(fname, "w") as f:
        print("People with suspect email addresses:", file=f)
        print(timestamp,  file=f)
        count=0
        for personname, person in gPersons.items():
            if len(person.Email) > 0:
                if "," in person.Email or " " in person.Email:
                    count+=1
                    print(f"   {personname} has a email address containing a comma or a space", file=f)
                else:
                    pattern=r"^[a-zA-Z0-9_]+@[a-zA-Z0-9_]+\.[a-zA-Z0-9]+$"
                    m=RegEx.match(pattern, person.Email)
                    if m is None:
                        count+=1
                        print(f"   {personname} has a email address not of the form something@something.something", file=f)

        if count == 0:
            print("    None found", file=f)


    #******
    # Check for people in the schedule whose response is 'y', but who are not scheduled to be on the program
    fname=os.path.join(reportsdir, "Diag - People response is 'y' but who are not scheduled.txt")
    with open(fname, "w") as f:
        print("People who are scheduled and in People but whose response is 'y' but who are not scheduled:", file=f)
        print(timestamp,  file=f)
        count=0
        for personname in gPersons.keys():
            if gPersons[personname].RespondedYes:
                found=False
                for item in gSchedules.values():
                    for x in item:
                        if personname == x.PersonName:
                            found=True
                            break
                    if found:
                        break
                if not found:
                    count+=1
                    print(f"   {personname} is not scheduled", file=f)
        if count == 0:
            print("    None found", file=f)


    # Does (t1, l1) overlap (t2, l2) where t and l and times and lengths in float hours?
    def TimesOverlap(t1: NumericTime, l1: float, t2: NumericTime, l2: float) -> bool:
        # Define epsilon=0.001 hours slop
        epsilon=0.001

        # Bogus times never overlap
        if t1.Bogus or t2.Bogus:
            return False

        # Note that we want to ignore 0-length overlaps such as  (10.0, 1.0) not overlapping (11.0, x)
        if t1 < t2:
            if t1+l1 > t2+epsilon:
                return True    # t1+l1 is less than t2 or exceeds t2 by less than epsilon
            return False
        # So t1 must be >= t2
        if t2+l2 > t1+epsilon:
            return True    # t2+l2 is less than t1 or exceeds t1 by less than epsilon
        return False


    # #******
    # # Check for people who are scheduled opposite themselves
    fname=os.path.join(reportsdir, "Diag - People with schedule conflicts.txt")
    with open(fname, "w") as f:
        print("People with schedule conflicts", file=f)
        print(timestamp,  file=f)
        count=0
        for personname in gSchedules.keys():
            pSched=[x for x in gSchedules[personname] if not x.IsDummy]     # Get a single person's schedule w/o dummy entries
            if len(pSched) == 0:
                continue

            # Look for duplicate times
            if len(pSched) > 1:     # Need two to tango
                pSched.sort(key=lambda x: x.Time)       # Sort pSched by time
                prev: ScheduleElement=pSched[0]
                for item in pSched[1:]:
                    # We insert dummy items for use elsewhere and need to ignore them here.  Also, prev is initialized to an empty Item which also has IsDummy set
                    if TimesOverlap(item.Time, item.Length, prev.Time, prev.Length):
                        print(f"{personname}: is scheduled to be in {prev.Room} and also {item.Room} at {prev.Time}", file=f)
                        count+=1
                    prev=item

            # Now check for Avoid conflicts
            avoidments=gPersons[personname].Avoid
            for item in pSched:
                for av in avoidments:
                    if TimesOverlap(item.Time, item.Length, av.Start, av.Duration):
                        print(f'{personname}: is scheduled to be in {item.Room} at {item.Time}, conflicting with "{av}"', file=f)
                        count+=1

        # To make it clear that the test ran, write a message if no conflicts were found.
        if count == 0:
            print("    None found", file=f)

    #******
    # Make a handy-dandy list of people's scheduling limitations
    fname=os.path.join(reportsdir, "People's scheduling limitations.txt")
    with open(fname, "w") as f:
        print("People's scheduling limitations", file=f)
        print(timestamp,  file=f)
        for personname in gSchedules.keys():
            avoidments=gPersons[personname].Avoid
            output=f"{personname}: "
            found=False
            for av in avoidments:
                if not found:
                    found=True
                else:
                    output+=", "
                output+=av.Pretty()
            if found:
                print(output, file=f)



    #******
    # Now look for similar name pairs
    # First we make up a list of all names that appear in any tab
    names=set()
    names.update(gSchedules.keys())
    names.update(gPersons.keys())
    similarNames: list[tuple[str, str, float]]=[]
    for p1 in names:
        for p2 in names:
            if p1 < p2:
                rat=difflib.SequenceMatcher(a=p1, b=p2).ratio()
                if rat > .75:
                    similarNames.append((p1, p2, rat))
    similarNames.sort(key=lambda x: x[2], reverse=True)

    fname=os.path.join(reportsdir, "Diag - Disturbingly similar names.txt")
    SafeDelete(fname)
    if len(similarNames) > 0:
        with open(fname, "w") as f:
            print("Names that are disturbingly similar:", file=f)
            print(timestamp,  file=f)
            count=0
            for s in similarNames:
                print(f"   {s[0]}  &  {s[1]}", file=f)
                count+=1
            if count == 0:
                print("    None found", file=f)


    # *********************************************************************************************************
    # *********************************************************************************************************
    # Now do the content/working reports

    #*******
    # Print the People with items by time report
    # Get a list of the program participants (the keys of the participants dictionary) sorted by the last token in the name (which will usually be the last name)
    sortedAllParticipantList=sorted(gSchedules.keys(), key=lambda x: x.split(" ")[-1])
    fname=os.path.join(reportsdir, "People with items by time.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("People with Items by Time\n", file=f)
        print(timestamp,  file=f)
        for personname in sortedAllParticipantList:
            if gPersons[personname].RespondedYes:
                print("\n"+personname, file=f)
                for schedElement in gSchedules[personname]:
                    if len(schedElement.DisplayName) > 0:
                        print(f"    {schedElement.Time}: {schedElement.DisplayName} [{schedElement.Room}] {schedElement.ModFlag}", file=f)

    #*******
    # Print the Items with people by time report
    # Get a list of the program participants (the keys of the  participants dictionary) sorted by the last token in the name (which will usually be the last name)
    fname=os.path.join(reportsdir, "Items with people by time.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("Items with People by Time\n", file=f)
        print(timestamp,  file=f)
        for time in gTimes:
            for room in gRoomNames:
                # Now search for the program item and people list for this slot
                for itemName, item in gItems.items():
                    if item.Time == time and item.Room == room:
                        print(f"{time}, {room}: {itemName}   {item.DisplayPlist()}", file=f)
                        if item.Precis is not None and item.Precis != "":
                            print("     "+ScrubPrecis(item.Precis), file=f)


    #*******
    # Print the program participant's schedule report in .txt and docx formats, simultaneously.
    # We print the .txt file as we go along, while accumulating the docx file in docx.Document() object, and then output it at the end.
    # Get a list of the program participants (the keys of the  participants dictionary) sorted by the last token in the name (which will usually be the last name)
    doc=docx.Document()     # The object holding the partly-created Word document
    fname=os.path.join(reportsdir, "Program participant schedules.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print(timestamp,  file=f)
        for personname in sortedAllParticipantList:
            if PersonOfInterest(personname, gSchedules):
                section=doc.add_section()
                section.orientation=WD_ORIENTATION.PORTRAIT
                print("\n\n********************************************", file=f)
                print(personname, file=f)
                AppendParaToDoc(doc, personname, bold=True, size=16)
                for schedElement in gSchedules[personname]:
                    if len(schedElement.DisplayName) > 0:
                        print(f"\n{schedElement.Time}: {schedElement.DisplayName} [{schedElement.Room}]", file=f)
                        para=doc.add_paragraph()
                        AppendTextToPara(para, f"\n{schedElement.Time}:", size=14)
                        AppendTextToPara(para, "  "+schedElement.DisplayName, size=14, bold=True)
                        AppendTextToPara(para, "  "+schedElement.Room, size=12)
                        item=gItems[schedElement.ItemName]
                        part=f"Participants: {item.DisplayPlist()}"
                        print(part, file=f)
                        AppendParaToDoc(doc, part)
                        if item.Precis is not None and item.Precis != "":
                            print(f"Precis: {item.Precis}", file=f)
                            AppendParaToDoc(doc, item.Precis, italic=True, size=12)
    # The .txt file has been written and closed, so now output the docx.Document() object as a Word file.
    fname=os.path.join(reportsdir, "Program participant schedules.docx")
    doc.save(fname)


    # *******
    # Print the program participant's schedule report
    # Get a list of the program participants (the keys of the  participants dictionary) sorted by the last token in the name (which will usually be the last name)
    fname=os.path.join(reportsdir, "Program participant schedules.xml")
    SafeDelete(fname)
    with open(fname, "w") as xml:
        for personname in sortedAllParticipantList:
            print(f"<person><full name>{personname}</full name>", file=xml)
            print(f"<email>{gPersons[personname].Email}</email>", file=xml)
            if sum(not x.IsDummy for x in gSchedules[personname]) == 0:
                print(f"<item><title>No Items Scheduled Yet</title><participants>{personname}</participants></item>", file=xml)
            else:
                for schedElement in gSchedules[personname]:
                    if len(schedElement.DisplayName) > 0:
                        print(f"<item><title>{schedElement.Time}: {schedElement.DisplayName} [{schedElement.Room}]</title>", file=xml)
                        item=gItems[schedElement.ItemName]
                        if schedElement.DisplayName in gItems and gItems[schedElement.DisplayName].Parms.Exists("equipment"):
                            print(f"<equipment>{gItems[schedElement.DisplayName].Parms['equipment']}</equipment>", file=xml)
                        print(f"<participants>{item.DisplayPlist()}</participants>", file=xml)
                        if item.Precis is not None and item.Precis != "":
                            print(f"<precis>{item.Precis}</precis>", file=xml)
                        print(f"</item>\n", file=xml)
            print("</person>", file=xml)


    #*******
    # Put out the entire People table in pseudo-XML format
    fname=os.path.join(reportsdir, "Program participants.xml")
    SafeDelete(fname)
    with open(fname, "w") as xml:
        for person in gPersons.values():
            xml.writelines(f"<person>")
            for key, val in person.Parms.items():
                xml.writelines(f"<{key}>{val}</{key}>")
            xml.writelines("</person>\n")


    #******
    # Report on the number of people/item
    fname=os.path.join(reportsdir, "Items' people counts.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        itemdata=[]
        for itemname, item in gItems.items():
            itemdata.append([len(item.People), str(item.Time), item.Name])
            print(f"{item.Time} {item.Name}: {len(item.People)}", file=f)

    fname=os.path.join(reportsdir, "Items' people counts.csv")
    with open(fname, mode='w', encoding='UTF8', newline="") as f:
        writer=csv.writer(f, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        writer.writerow(["Number", "Item Time", "Item Title"])
        for id in itemdata:
            writer.writerow(id)


    #******
    # Flag items with a suspiciously small number of people on them
    fname=os.path.join(reportsdir, "Diag - Items with unexpectedly low number of participants.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("List of non-readings, non-KKs, and non-solo items with fewer than 3 people on them\n\n", file=f)
        print(timestamp,  file=f)
        found=False
        for itemname, item in gItems.items():
            if item.Name:
                if len(item.People) >= 3:
                    continue
                if "Reading" in item.Name or "KK" in item.Name or "Kaffe" in item.Name or "Autograph" in item.Name:
                    continue
                if item.Parms["solo"]:
                    continue
                print(f"{item.Time} {item.Name}: {len(item.People)}", file=f)
                found=True
        if not found:
            print("None found", file=f)


    #******
    # Flag items missing a moderator or a precis
    fname=os.path.join(reportsdir, "Diag - Items missing a moderator.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("List of non-readings and KKs with no moderator\n\n", file=f)
        print(timestamp,  file=f)
        found=False
        for itemname, item in gItems.items():
            if "Reading" in item.Name or "KK" in item.Name or "Kaffe" in item.Name or "Autograph" in item.Name:
                continue
            if item.Parms["solo"]:  # Solo items don't need a moderator
                continue
            if item.ModName != "":
                continue
            print(f"{item.Time} {item.Name}: {len(item.People)}", file=f)
            found=True
        if not found:
            print("None found", file=f)


    fname=os.path.join(reportsdir, "Diag - Items missing a precis.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("List of non-readings and KKs with no precis\n\n", file=f)
        print(timestamp,  file=f)
        found=False
        for itemname, item in gItems.items():
            if "Reading" in item.Name or "KK" in item.Name or "Kaffe" in item.Name or "Autograph" in item.Name:
                continue
            if item.Precis is not None and len(item.Precis) > 0:
                continue
            print(f"{item.Time} {item.Name}: {len(item.People)}", file=f)
            found=True
        if not found:
            print("None found", file=f)



    fname=os.path.join(reportsdir, "Equipment requirements.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("List of items with equipment requirements\n\n", file=f)
        print(timestamp,  file=f)
        found=False
        for itemname, item in gItems.items():
            if item.Parms.Exists("equipment"):
                print(f"{item.Time}, {item.Room}:  {item.Name}\n\t\t{item.Parms['equipment']}\n", file=f)
                found=True
        if not found:
            print("None found", file=f)


    #******
    # Report on the number of items/person
    # Include all people in the people tab, even those with no items
    fname=os.path.join(reportsdir, "Peoples' item counts.txt")
    SafeDelete(fname)
    with open(fname, "w") as f:
        print("List of number of items each person is scheduled on\n", file=f)
        print(timestamp,  file=f)
        for personname, person in gPersons.items():
            if PersonOfInterest(person, gSchedules):
                if personname in gSchedules.keys():
                    numItems=sum(not x.IsDummy for x in gSchedules[personname])
                    print(f"{personname}: {numItems}{'' if person.RespondedYes else ' not confirmed'}", file=f)
                else:
                    if person.RespondedYes:
                        print(personname+": responded Yes, but is not scheduled", file=f)

    fname=os.path.join(reportsdir, "Peoples' item counts.csv")
    with open(fname, "w", encoding='UTF8', newline="") as f:
        writer=csv.writer(f, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
        writer.writerow(["Number" , "Person"])
        for personname, person in gPersons.items():
            if PersonOfInterest(person, gSchedules):
                numItems=sum(not x.IsDummy for x in gSchedules[personname])
                writer.writerow([numItems, personname])


    # Create the pocket program Word file and the .txt file at the same time
    fname=os.path.join(reportsdir, "Pocket program.txt")
    try:
        if not SafeDelete(fname):
            pass
    except:
        pass
    try:
        f=open(fname, "w")    # The file to receive the .txt document
    except:
        pass
        # Popup("open("+fname+")  threw exception")

    doc=docx.Document()     # The object holding the partly-created Word document
    AppendParaToDoc(doc, "Schedule", bold=True, size=24)
    print("Schedule", file=f)
    for time in gTimes:
        AppendParaToDoc(doc, "")
        AppendParaToDoc(doc, str(time), bold=True)
        print(f"\n{time}", file=f)
        for room in gRoomNames:
            # Now search for the program item and people list for this slot
            for itemName, item in gItems.items():
                if item.Time == time and item.Room == room:
                    if len(item.DisplayName) > 0:
                        para=doc.add_paragraph()
                        AppendTextToPara(para, room+": ", italic=True, size=12, indent=0.3)
                        AppendTextToPara(para, item.DisplayName, size=12, indent=0.3)
                        print(f"   {room}:  {item.DisplayName}", file=f)   # Print the room and item name
                        if len(item.People) > 0:            # And the item's people list
                            plist=item.DisplayPlist()
                            AppendParaToDoc(doc, plist, size=12, indent=0.6)
                            print("            "+plist, file=f)
                        if item.Precis is not None and item.Precis != "":
                            AppendParaToDoc(doc, ScrubPrecis(item.Precis), italic=True, size=12, indent=0.6)
                            print("            "+ScrubPrecis(item.Precis), file=f)
    fname=os.path.join(reportsdir, "Pocket program.docx")
    doc.save(fname)
    f.close()


    # Create the individual (one per person) tentcard Word document
    doc=docx.Document()
    for personname in sortedAllParticipantList:
        if any([not x.IsDummy for x in gSchedules[personname]]):
            section=doc.add_section()
            section.orientation=WD_ORIENTATION.LANDSCAPE
            section.page_width=Inches(11)
            section.page_height=Inches(8.5)
            section.top_margin=Inches(5)
            section.bottom_margin=Inches(1)
            section.right_margin=Inches(0.2)
            section.left_margin=Inches(0.2)

            para=doc.add_paragraph()
            para.alignment=1
            size=86
            if len(personname) > 18:
                size=86*18/len(personname)
            AppendTextToPara(para, personname, size=size, indent=0)

    doc.save(os.path.join(reportsdir, "Tentcards -- Individual.docx"))


    # Create the tentcards for each program item Word document
    doc=docx.Document()
    for room in gRoomNames:
        for time in gTimes:
            for itemName, item in gItems.items():
                if item.Time == time and item.Room == room:
                    if len(item.DisplayName) > 0:
                        for person in item.People:
                            # Do a tentcard for this person
                            section=doc.add_section()
                            section.orientation=WD_ORIENTATION.LANDSCAPE
                            section.page_width=Inches(11)
                            section.page_height=Inches(8.5)

                            section.top_margin=Inches(1)
                            section.right_margin=Inches(0.2)
                            section.left_margin=Inches(0.2)
                            #section.top_margin=Inches(5)
                            section.bottom_margin=Inches(1)

                            # Add the paragraph for this tentcard
                            AppendParaToDoc(doc, f"{time} --  {room}\n", size=12, indent=0)
                            AppendParaToDoc(doc, f"{item.DisplayName}\n", size=12, indent=0)

                            # Set the margins for the big person's name for the front of the tentcard
                            AppendParaToDoc(doc, "\n", size=230)
                            size=86
                            if len(person) > 18:
                                size=86*18/len(person)
                            AppendParaToDoc(doc, person, size=size, indent=0, alignment=1)

    doc.save(os.path.join(reportsdir, "Tentcards -- By Program Item.docx"))


    #******
    # Generate web pages, one for each day.
    currentday=""
    f=None
    for time in gTimes:
        # We generate a separate report for each day
        # The times are sorted in ascending order.
        # We will let the act of the time flipping over to a new day create the new file
        sortday=time.NominalDayString
        if sortday != currentday:
            # Close the old file, if any
            if f is not None:
                f.write('</font></table>\n')
                # Read and append the footer
                try:
                    with open(PyiResourcePath("control-WebpageFooter.txt")) as f2:
                        f.writelines(f2.readlines())
                except:
                    MessageLog("Can't read 'control-WebpageFooter.txt' (1)")
                f.close()
                f=None
            # And open the new file
            currentday=sortday
            fname=os.path.join(reportsdir, "Schedule - "+sortday+".html")
            SafeDelete(fname)
            f=open(fname, "w")
            try:
                with open(PyiResourcePath("control-WebpageHeader.txt")) as f2:
                    try:
                        f.writelines(f2.readlines())
                    except:
                        LogError("Failure copying 'control-WebpageHeader.txt'")
            except:
                MessageLog("Can't open 'control-WebpageHeader.txt'")
            f.write("<h2>"+sortday+"</h2>\n")
            f.write('<table border="0" cellspacing="0" cellpadding="2">\n')

        f.write('<tr><td colspan="3">')
        f.write(f'<p class="time">{time.NumericToTextTime()}</p>')
        f.write('</td></tr>\n')
        for room in gRoomNames:
            # Now search for the program item and people list for this slot
            for itemName, item in gItems.items():
                if len(item.DisplayName) > 0:
                    if item.Time == time and item.Room == room:
                        f.write('<tr><td width="40">&nbsp;</td><td colspan="2">')   # Two columns, the first 40 pixes wide and empty
                        f.write(f'<p><span class="room">{room}: </span><span class="item">{item.DisplayName}</span></p>')
                        f.write('</td></tr>')
                        if len(item.People) > 0:            # And the item's people list
                            f.write('<tr><td width="40">&nbsp;</td><td width="40">&nbsp;</td><td width="600">')     # Three columns, the first two 40 pixes wide and empty; the third 600 pixels wide
                            f.write(f'<p><span class="people">{UnicodeToHtml(item.DisplayPlist())}</span></p>')
                            f.write('</td></tr>\n')
                        if item.Precis is not None and item.Precis != "":
                            f.write('<tr><td width="40">&nbsp;</td><td width="40">&nbsp;</td><td width="600">')     # Same
                            f.write(f'<p><span class="precis">{UnicodeToHtml(ScrubPrecis(item.Precis))}</span></p>')
                            f.write('</td></tr>\n')
    if f is not None:
        # Read and append the footer
        f.write('</table>\n')
        try:
            with open(PyiResourcePath("control-WebpageFooter.txt")) as f2:
                f.writelines(f2.readlines())
        except:
            MessageLog("Can't read 'control-WebpageFooter.txt' (2)")
        f.close()


    #******
    # Do the room signs.  They'll go in reports/rooms/<name>.docx
    # Create the roomsigns subfolder if none exists
    path=os.path.join(reportsdir, "roomsigns")
    if not os.path.exists(path):
        os.mkdir(path)
    for room in gRoomNames:
        inuse=False  # Make sure that this room is actually in use
        if len(room.strip()) == 0:
            continue
        doc=docx.Document()
        AppendParaToDoc(doc, room, bold=True, size=32)  # Room name at top
        for time in gTimes:
            for itemName in gItems.keys():
                item=gItems[itemName]
                if len(item.DisplayName) > 0:
                    if item.Time == time and item.Room == room:
                        inuse=True
                        AppendParaToDoc(doc, "")    # Skip a line
                        para=doc.add_paragraph()
                        AppendTextToPara(para, f"{item.Time}:  ", bold=True)   # Add the time in bold followed by the item's title
                        AppendTextToPara(para, item.DisplayName)
                        AppendParaToDoc(doc, item.DisplayPlist(), italic=True, indent=0.5)        # Then, on a new line, the people list in italic
        fname=os.path.join(path, room.replace("/", "-")+".docx")
        SafeDelete(fname)
        if inuse:
            doc.save(fname)

    Log(f"Reports generated in directory '{reportsdir}'")
    Log("Done.")
    LogClose()


#*************************************************************************************************
#*************************************************************************************************
# Miscellaneous helper functions

# Read the contents of a Google docs spreadsheet tab into a list of lists of strings
# Ignore rows beginning with #
def ReadSheetFromGoogleTab(sheet, spreadSheetID, parms: ParmDict, parmname: str) -> list[list[str]]:

    # Convert the generic name of the tab to the specific name to be used this year
    tabname=GetParmFromParmDict(parms, parmname)
    try:
        cells=sheet.values().get(spreadsheetId=spreadSheetID, range=f'{tabname}!A1:Z999').execute().get('values', [])  # Read the whole thing.
    except HttpError:
        LogError(f"ReadSheetFromTab: Can't locate {tabname} tab in spreadsheet. Is the supplied SheetID wrong?")
        exit(999)
    except Exception as e:
        LogError(f"ReadSheetFromTab: Exception {e} while attempting to load tab {tabname} tab in spreadsheet.")
        exit(999)

    if not cells:
        LogError(f"ReadSheetFromTab: Can't locate {tabname} tab in spreadsheet")
        raise (ValueError, "No cells found in tab")

    rows=[p for p in cells if len(p) > 0 and "".join(p)[0] != "#"]  # Drop empty lines and lines with a "#" alone in column 1.
    return rows

def ReadSheetFromXLSXTab(workbook: openpyxl.Workbook, parms: ParmDict, parmname: str) -> list[list[str]]:

    # Convert the generic name of the tab to the specific name to be used this year
    tabname=GetParmFromParmDict(parms, parmname)
    if tabname not in workbook.sheetnames:
        LogError(f"ReadSheetFromTab: Can't locate {tabname} tab in spreadsheet")
        raise (ValueError, f"No cells found in tab '{tabname}'")

    rows=workbook[tabname].values
    rows=[list(row) for row in rows]        # Turn rows (supplied as tuples by values) into lists
    rows=[["" if cell is None else cell for cell in row] for row in rows]   # Turn None values into empty strings
    rows=[row for row in rows if any([cell != "" for cell in row])]         # Eliminate entirely empty rows
    rows=[[str(cell) for cell in row] for row in rows]              # Some cells seem to come through as ints -- turn them into strs
    rows=[row for row in rows if "".join(row)[0] != "#"]            # Ignore rows where the 1st character is a "#"
    trimmed=[]      # Remove trailing empty cells (Probably not needed, but better duplicates what Googledocs returns.)
    for row in rows:
        while row[-1] == "":
            row.pop()
        trimmed.append(row)
    return trimmed


# Take a Person and gSchedules and return True if that Person is scheduled on some item *or* is listed as Response='y'
def PersonOfInterest(person: [Person, str], gschedules: dict[str, list[ScheduleElement]]) -> bool:
    if type(person) is Person:
        # These checks only apply for a Person structure
        if person.RespondedYes:
            return True
        if person.Fullname not in gschedules.keys():
            return False
        # The last test needs person to be a string
        person=person.Fullname
        
    return sum(not x.IsDummy for x in gschedules[person]) > 0


# Take a name string which may contain the (M) moderater flag and split it into isMon and the name by itself
# Generate the name of a person stripped if any "(M)" or "(m)" flags
def CheckModFlag(s: str) -> tuple[bool, str]:
    if "(m)" in s.lower():
        return True, s.replace("(M)", "").replace("(m)", "").strip()
    return False, s


# Delete a file, ignoring any errors
# We do this because of as-yet not understood failures to delete files
def SafeDelete(fn: str) -> bool:
    try:
        os.remove(fn)
    except:
        return False
    return True

#.......
# Add an item with a list of people to the gItems dict, and add the item to each of the persons who are on it
def AddItemWithPeople(gItems: dict[str, Item], time: NumericTime, roomName: str, itemName: str, plistText: str, length: float=1.0) -> None:

    # Ignore anything following a "#" as a comment
    if "#" in plistText:
        plistText=plistText[:plistText.index("#")]
    plist=[p.strip() for p in plistText.split(",") ]    # Get the people as a list with excess spaces removed
    plist=[p for p in plist if len(p) > 0]              # Ignore empty entries
    modName=""
    peopleList: list[str]=[]
    for person in plist:  # For each person listed on this item
        ismod, name=CheckModFlag(person)
        if ismod:
            modName=name
        peopleList.append(name)
    # And add the item with its list of people to the items table.
    if itemName in gItems:  # If the item's name is already in use, add a uniquifier of room+day/time
        itemName='{'+f"{itemName}  {roomName} {time}"+'}'
    item=Item(ItemText=itemName, Time=time, Length=length, Room=roomName, People=peopleList, ModName=modName)
    gItems[item.Name]=item


#.......
# Add an item with a list of people, and add the item to each of the persons
def AddItemWithoutPeople(gItems: dict[str, Item], time: NumericTime, roomName: str, itemName: str, length: float=0) -> None:
    if itemName in gItems:  # If the item's name is already in use, add a uniquifier of room+day/time
        itemName=itemName+'  {'+f"{roomName} {time}"+'}'
    item=Item(ItemText=itemName, Time=time, Room=roomName, Length=length)
    gItems[item.Name]=item


#******
# Create a docx and a .txt version for the pocket program
def AppendParaToDoc(doc: docx.Document, txt: str, bold=False, italic=False, size=14, indent=0.0, font="Calibri", alignment=0):
    para=doc.add_paragraph()
    run=para.add_run(txt)
    run.bold=bold
    run.italic=italic
    runfont=run.font
    runfont.name=font
    runfont.size=Pt(size)
    para.alignment=alignment
    para.paragraph_format.left_indent=Inches(indent)
    para.paragraph_format.line_spacing=1
    para.paragraph_format.space_after=0

def AppendTextToPara(para: docx.text.paragraph.Paragraph, txt: str, bold: bool=False, italic: bool=False, size: float=14, indent: float=0.0, font: str="Calibri"):
    run=para.add_run(txt)
    run.bold=bold
    run.italic=italic
    runfont=run.font
    runfont.name=font
    runfont.size=Pt(size)
    para.paragraph_format.left_indent=Inches(indent)
    para.paragraph_format.line_spacing=1
    para.paragraph_format.space_after=0



if __name__ == "__main__":
    main()